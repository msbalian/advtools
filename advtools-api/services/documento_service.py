import os
import asyncio
import uuid
import shutil
import re
import unicodedata
import json
from datetime import datetime
import locale
import locale
from fastapi import HTTPException, UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from docx import Document # Importando Document para leitura direta do texto
import schemas
import models
import crud
from docxtpl import DocxTemplate
from services.storage_service import get_storage_provider
from services.ai_service import redigir_documento_com_ia
from config import Config

def get_docx_text(path: str) -> str:
    """Extrai o texto puro de um documento .docx para servir de contexto à IA."""
    try:
        if not os.path.exists(path):
            return ""
        doc = Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Erro ao ler texto do DOCX: {e}")
        return ""

try:
    locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')
except Exception:
    try:
        locale.setlocale(locale.LC_TIME, 'pt-BR')
    except Exception:
        try:
           locale.setlocale(locale.LC_TIME, 'pt_BR.utf8')
        except Exception:
           pass

async def convert_docx_to_pdf_async(in_path: str, out_path: str):
    """Executa a conversão do DOCX para PDF em uma thread separada."""
    abs_in = os.path.abspath(in_path)
    abs_out = os.path.abspath(out_path)
    
    loop = asyncio.get_running_loop()
    
    def do_convert():
        import pythoncom
        from docx2pdf import convert
        pythoncom.CoInitialize()
        try:
            convert(abs_in, abs_out)
        finally:
            pythoncom.CoUninitialize()

    await loop.run_in_executor(None, do_convert)

# --- ROTAS DE MODELOS ---

async def listar_modelos_service(db: AsyncSession, escritorio_id: int):
    q = select(models.ModeloDocumento).where(models.ModeloDocumento.escritorio_id == escritorio_id)
    res = await db.execute(q)
    modelos = res.scalars().all()
    if not modelos:
        return []
    return [{"id": m.id, "nome": m.nome, "arquivo_path": m.arquivo_path, "data_criacao": m.data_criacao} for m in modelos]

async def criar_modelo_service(db: AsyncSession, current_user: models.Usuario, nome: str, file: UploadFile):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Apenas arquivos .docx são permitidos")
        
    escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
    storage = get_storage_provider(escritorio)
    
    content = await file.read()
    unique_filename = f"{uuid.uuid4().hex}_{file.filename}"
    
    # Modelos ficam em uma subpasta 'modelos' dentro da raiz do escritório/local
    db_path = await storage.save_file(content, "modelos", unique_filename)
        
    novo_modelo = models.ModeloDocumento(
        escritorio_id=current_user.escritorio_id,
        nome=nome,
        arquivo_path=db_path
    )
    db.add(novo_modelo)
    await db.commit()
    await db.refresh(novo_modelo)
    
    return {"id": novo_modelo.id, "nome": novo_modelo.nome, "arquivo_path": novo_modelo.arquivo_path}

async def deletar_modelo_service(db: AsyncSession, current_user: models.Usuario, modelo_id: int):
    q = select(models.ModeloDocumento).where(
        models.ModeloDocumento.id == modelo_id,
        models.ModeloDocumento.escritorio_id == current_user.escritorio_id
    )
    res = await db.execute(q)
    modelo = res.scalars().first()
    
    if not modelo:
        raise HTTPException(status_code=404, detail="Modelo não encontrado")
        
    escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
    storage = get_storage_provider(escritorio)
    await storage.delete_file(modelo.arquivo_path)
        
    await db.delete(modelo)
    await db.commit()
    return {"detail": "Modelo deletado com sucesso"}

async def substituir_modelo_service(db: AsyncSession, current_user: models.Usuario, modelo_id: int, file: UploadFile):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Apenas arquivos .docx são permitidos")
        
    q = select(models.ModeloDocumento).where(
        models.ModeloDocumento.id == modelo_id,
        models.ModeloDocumento.escritorio_id == current_user.escritorio_id
    )
    res = await db.execute(q)
    modelo = res.scalars().first()
    
    if not modelo:
        raise HTTPException(status_code=404, detail="Modelo não encontrado")
        
    escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
    storage = get_storage_provider(escritorio)
    
    # Remove antigo
    await storage.delete_file(modelo.arquivo_path)
    
    # Salva novo
    content = await file.read()
    unique_filename = f"{uuid.uuid4().hex}_{file.filename}"
    db_path = await storage.save_file(content, "modelos", unique_filename)
        
    modelo.arquivo_path = db_path
    await db.commit()
    await db.refresh(modelo)
    
    return {"id": modelo.id, "nome": modelo.nome, "arquivo_path": modelo.arquivo_path}

# --- ROTAS DE DOCUMENTOS DO CLIENTE ---

async def read_documentos_cliente_service(db: AsyncSession, current_user: models.Usuario, cliente_id: int, pasta_id: int = -1):
    return await crud.get_documentos_cliente(db, cliente_id, escritorio_id=current_user.escritorio_id, pasta_id=pasta_id)

async def read_documentos_escritorio_service(db: AsyncSession, current_user: models.Usuario, pasta_id: int = -1):
    return await crud.get_documentos_escritorio(db, escritorio_id=current_user.escritorio_id, pasta_id=pasta_id)

async def read_documento_service(db: AsyncSession, current_user: models.Usuario, documento_id: int):
    doc = await crud.get_documento_by_id(db, documento_id, current_user.escritorio_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    return doc

async def upload_documento_cliente_service(db: AsyncSession, current_user: models.Usuario, cliente_id: int, file: UploadFile, nome: str, **kwargs):
    escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
    storage = get_storage_provider(escritorio)
    
    content = await file.read()
    unique_filename = f"{uuid.uuid4().hex}_{file.filename}"
    
    # Organização: cliente_{id}/documentos
    relative_dir = f"cliente_{cliente_id}/documentos"
    db_path = await storage.save_file(content, relative_dir, unique_filename)
    
    doc_create = schemas.DocumentoClienteCreate(nome=nome, cliente_id=cliente_id, pasta_id=kwargs.get('pasta_id'))
    return await crud.create_documento_cliente(
        db, 
        doc_create, 
        db_path, 
        current_user.escritorio_id
    )

async def upload_documento_escritorio_service(db: AsyncSession, current_user: models.Usuario, file: UploadFile, nome: str, pasta_id: int = -1):
    escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
    storage = get_storage_provider(escritorio)
    
    content = await file.read()
    unique_filename = f"{uuid.uuid4().hex}_{file.filename}"
    
    # Organização: escritorio/documentos
    relative_dir = "escritorio/documentos"
    db_path = await storage.save_file(content, relative_dir, unique_filename)
    
    doc_create = schemas.DocumentoClienteCreate(nome=nome, cliente_id=None, pasta_id=pasta_id if pasta_id != -1 else None)
    return await crud.create_documento_cliente(
        db, 
        doc_create, 
        db_path, 
        current_user.escritorio_id
    )

async def update_documento_file_service(db: AsyncSession, current_user: models.Usuario, documento_id: int, file: UploadFile):
    doc = await crud.get_documento_by_id(db, documento_id, current_user.escritorio_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
        
    escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
    storage = get_storage_provider(escritorio)
    
    content = await file.read()
    # Atualiza mantendo a organização
    if doc.cliente_id:
        relative_dir = f"cliente_{doc.cliente_id}/documentos"
    else:
        relative_dir = "escritorio/documentos"
        
    filename = os.path.basename(doc.arquivo_path) if doc.arquivo_path else f"{uuid.uuid4().hex}_{file.filename}"
    
    db_path = await storage.save_file(content, relative_dir, filename)
    doc.arquivo_path = db_path
        
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc

async def delete_documento_service(db: AsyncSession, current_user: models.Usuario, documento_id: int):
    deleted_doc = await crud.delete_documento_cliente(db, documento_id, current_user.escritorio_id)
    if not deleted_doc:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
        
    file_path = os.path.join("static", deleted_doc.arquivo_path)
    if not os.path.exists(file_path):
        file_path = os.path.join("static/armazenamento", deleted_doc.arquivo_path)
        
    if os.path.exists(file_path):
        os.remove(file_path)
    
    return True

# --- REDATOR INTELIGENTE ---

def slugify(value):
    nfkd_form = unicodedata.normalize('NFKD', value)
    only_ascii = nfkd_form.encode('ascii', 'ignore').decode('ascii')
    cleaned = re.sub(r'[^\w\s-]', '', only_ascii).strip().lower()
    return re.sub(r'[-\s]+', '_', cleaned)

async def gerar_documento_service(db: AsyncSession, current_user: models.Usuario, request: schemas.GerarDocumentoRequest):
    cliente = None
    if request.cliente_id:
        cliente = await crud.get_cliente(db, request.cliente_id, current_user.escritorio_id)
        if not cliente:
            raise HTTPException(status_code=404, detail="Cliente não encontrado")
        
    q_modelo = select(models.ModeloDocumento).where(
        models.ModeloDocumento.id == request.modelo_id,
        models.ModeloDocumento.escritorio_id == current_user.escritorio_id
    )
    res_modelo = await db.execute(q_modelo)
    modelo = res_modelo.scalars().first()
    
    if not modelo:
        raise HTTPException(status_code=404, detail="Modelo não encontrado")

    source_path = os.path.join("static", modelo.arquivo_path)
    if not os.path.exists(source_path):
        source_path = os.path.join("static/armazenamento", modelo.arquivo_path)

    partes = []
    servicos = []
    if request.cliente_id:
        partes = await crud.get_partes_cliente(db, request.cliente_id, current_user.escritorio_id)
        servicos = await crud.get_servicos_by_cliente(db, request.cliente_id, current_user.escritorio_id)
    
    servico = next((s for s in servicos if s.status == "Ativo"), servicos[0] if servicos else None)

    hoje = datetime.now()
    esc_obj = await crud.get_escritorio(db, current_user.escritorio_id)
    
    context = {
        "DADOS_DO_ESCRITORIO": {
            "nome": esc_obj.nome or "",
            "documento": esc_obj.documento or "",
        },
        "DADOS_DO_CLIENTE_PRINCIPAL": {
            "nome": cliente.nome if cliente else esc_obj.nome or "",
            "documento": cliente.documento if cliente else esc_obj.documento or "",
            "endereco_completo": (f"{cliente.endereco}, {cliente.bairro}, {cliente.cidade}-{cliente.uf}, CEP: {cliente.cep}") if cliente else "",
            "email": (cliente.email or "") if cliente else "",
            "nacionalidade": (cliente.nacionalidade or "") if cliente else "",
            "estado_civil": (cliente.estado_civil or "") if cliente else "",
            "profissao": (cliente.profissao or "") if cliente else "",
            "rg": (cliente.rg or "") if cliente else "",
            "data_nascimento": (cliente.data_nascimento or "") if cliente else "",
        },
        "cliente_nome": (cliente.nome or "") if cliente else esc_obj.nome or "",
        "cliente_doc": (cliente.documento or "") if cliente else esc_obj.documento or "",
        "cliente_endereco": (cliente.endereco or "") if cliente else "",
        "cliente_bairro": (cliente.bairro or "") if cliente else "",
        "cliente_cidade": (cliente.cidade or "") if cliente else "",
        "cliente_uf": (cliente.uf or "") if cliente else "",
        "cliente_cep": (cliente.cep or "") if cliente else "",
        "cliente_email": (cliente.email or "") if cliente else "",
        "cliente_nacionalidade": (cliente.nacionalidade or "") if cliente else "",
        "cliente_estado_civil": (cliente.estado_civil or "") if cliente else "",
        "cliente_profissao": (cliente.profissao or "") if cliente else "",
        "cliente_rg": (cliente.rg or "") if cliente else "",
        "cliente_data_nascimento": (cliente.data_nascimento or "") if cliente else "",
        "data_hoje": hoje.strftime("%d/%m/%Y"),
        "ano_atual": hoje.strftime("%Y"),
        "data_extenso": hoje.strftime("%d de %B de %Y"),
        "conteudo_ia": "", 
        "servico_tipo": (servico.tipo_servico_id if servico else "") if cliente else "",
        "percentual_exito": (servico.porcentagem_exito if servico else "") if cliente else "",
    }

    # --- LÓGICA DE IA (REDAÇÃO INTEGRAL) ---
    if request.usar_ia:
        escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
        # Usa a chave do escritório ou a global do .env como fallback
        api_key = (escritorio.gemini_api_key if escritorio else None) or Config.GEMINI_API_KEY
        
        if api_key:
            # Extrair o TEXTO REAL do modelo para a IA entender o contexto
            modelo_full_text = get_docx_text(source_path)
            if not modelo_full_text:
                modelo_full_text = f"O advogado quer redigir um(a) {modelo.nome}."
                
            instrucoes = request.instrucoes_ia or "Redija o documento completo, adaptando o modelo fornecido aos dados do cliente e garantindo que toda a peça esteja juridicamente coerente."
            
            conteudo_gerado = await redigir_documento_com_ia(
                api_key=escritorio.gemini_api_key,
                modelo_texto=modelo_full_text,
                context=context,
                instrucoes=instrucoes
            )
            # Injetamos o resultado na tag de conteúdo. 
            # DICA: Se o modelo não tiver a tag, usaremos docxtpl.re_render ou fallback mais à frente.
            context["conteudo_ia"] = conteudo_gerado
        else:
            context["conteudo_ia"] = "[ERRO: Gemini API Key não configurada nas configurações do escritório.]"

    pagamentos_list = []
    pagamentos_resumo_linhas = []
    soma_valor_total = 0.0

    if servico and servico.condicoes_pagamento:
        try:
            pagamentos_raw = json.loads(servico.condicoes_pagamento)
            if isinstance(pagamentos_raw, list):
                pagamentos_list = pagamentos_raw
        except json.JSONDecodeError:
            pass
            
    for i, pag in enumerate(pagamentos_list):
        valor_num = float(pag.get("valor") or 0)
        soma_valor_total += valor_num
        valor_str = f"R$ {valor_num:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        
        data_str = pag.get("data") or ""
        if data_str and "T" not in data_str and len(data_str.split("-")) == 3:
            partes_data = data_str.split("-")
            data_str = f"{partes_data[2]}/{partes_data[1]}/{partes_data[0]}"

        tipo = pag.get("tipo") or "Desconhecido"
        obs = pag.get("obs") or ""
        
        linha_resumo = f"{i+1}ª Parcela - {data_str} - {valor_str} ({tipo})"
        if obs:
            linha_resumo += f" - {obs}"
        pagamentos_resumo_linhas.append(linha_resumo)
        
        if i < 12:
            prefix = f"pagamento_{i+1}_"
            context[f"{prefix}valor"] = valor_str
            context[f"{prefix}data"] = data_str
            context[f"{prefix}tipo"] = tipo
            context[f"{prefix}obs"] = obs

    for i in range(len(pagamentos_list), 12):
        prefix = f"pagamento_{i+1}_"
        context[f"{prefix}valor"] = ""
        context[f"{prefix}data"] = ""
        context[f"{prefix}tipo"] = ""
        context[f"{prefix}obs"] = ""

    context["pagamentos"] = pagamentos_list
    context["pagamentos_resumo"] = "\n".join(pagamentos_resumo_linhas)
    context["qtd_parcelas"] = str(len(pagamentos_list))
    
    parcelas_linhas = []
    for i, pag in enumerate(pagamentos_list):
        valor_num = float(pag.get("valor") or 0)
        valor_str = f"R$ {valor_num:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        
        data_str = pag.get("data") or ""
        if data_str and "T" not in data_str and len(data_str.split("-")) == 3:
            partes_data = data_str.split("-")
            data_str = f"{partes_data[2]}/{partes_data[1]}/{partes_data[0]}"

        linha = f"{i + 1}ª Parcela - {data_str} - {valor_str}"
        tipo = pag.get("tipo") or ""
        if tipo:
            linha += f" ({tipo})"
        obs = pag.get("obs") or ""
        if obs:
            linha += f" - {obs}"
            
        parcelas_linhas.append(linha)
        
    context["pagamento_parcelas"] = "\n".join(parcelas_linhas) if parcelas_linhas else ""
    
    if soma_valor_total > 0:
        context["pagamento_valor"] = f"R$ {soma_valor_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        context["valor_total"] = context["pagamento_valor"] 
    else:
        val = servico.valor_total if servico else 0.0
        context["pagamento_valor"] = f"R$ {val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") if val else ""
        context["valor_total"] = context["pagamento_valor"]

    context["forma_pagamento"] = servico.forma_pagamento if servico else ""
    context["detalhes_pagamento"] = servico.detalhes_pagamento if servico else ""
    
    partes_txt = []
    for i, parte in enumerate(partes):
        papel = parte.papel or "Envolvido"
        partes_txt.append(f"{parte.nome} ({papel})")
        
        if i < 5:
            prefix = f"parte_{i+1}_"
            context[f"{prefix}nome"] = parte.nome or ""
            context[f"{prefix}doc"] = parte.documento or ""
            context[f"{prefix}rg"] = parte.rg or ""
            context[f"{prefix}nasc"] = parte.data_nascimento or ""
            context[f"{prefix}nacionalidade"] = parte.nacionalidade or ""
            context[f"{prefix}estado_civil"] = parte.estado_civil or ""
            context[f"{prefix}profissao"] = parte.profissao or ""
            context[f"{prefix}endereco"] = parte.endereco or ""
            context[f"{prefix}bairro"] = parte.bairro or ""
            context[f"{prefix}cidade"] = parte.cidade or ""
            context[f"{prefix}uf"] = parte.uf or ""
            context[f"{prefix}cep"] = parte.cep or ""
            context[f"{prefix}papel"] = parte.papel or ""
            
    context["partes_resumo"] = "\n".join(partes_txt)
    context["partes_txt"] = ", ".join(partes_txt)

    for i in range(len(partes), 5):
        prefix = f"parte_{i+1}_"
        for field in ["nome", "doc", "rg", "nasc", "nacionalidade", "estado_civil", "profissao", "endereco", "bairro", "cidade", "uf", "cep", "papel"]:
            context[f"{prefix}{field}"] = ""

    try:
        # source_path já foi definido e verificado no início da função
        
        doc = DocxTemplate(source_path)
        doc.render(context)
        
        # --- VERIFICAÇÃO DE USO DE IA SEM TAG ---
        # Se usamos IA mas o documento resultou muito curto (provavelmente a tag {{conteudo_ia}} não existia)
        # OU se queremos forçar o conteúdo da IA como principal
        if request.usar_ia and context.get("conteudo_ia"):
            # Uma forma simples de detectar se a tag foi usada é ver se o texto da IA aparece no doc renderizado
            # Mas para garantir o desejo do usuário (IA adapta a peça), se a IA foi requisitada, 
            # e o documento final não parece ter mudado (ou apenas trocou tags básicas), 
            # podemos gerar um DOCX puramente com o texto da IA.
            
            # Decisão: Se usar_ia for True, vamos gerar um DOCX baseado no retorno da IA 
            # para garantir que a "adaptação" mencionada pelo usuário ocorra.
            new_doc = Document()
            
            # Tentar adicionar logo do escritório no topo se existir
            escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
            if escritorio.logo_path:
                logo_abs_path = os.path.join("static", escritorio.logo_path)
                if os.path.exists(logo_abs_path):
                    try:
                        new_doc.add_picture(logo_abs_path, width=None) 
                    except: pass
            
            # Adicionar o texto da IA
            for line in context["conteudo_ia"].split('\n'):
                if line.strip():
                    new_doc.add_paragraph(line)
                else:
                    new_doc.add_paragraph("")
            
            import io
            file_stream = io.BytesIO()
            new_doc.save(file_stream)
            file_content = file_stream.getvalue()
        else:
            # Fluxo padrão (tags ou IA com tag presente no docx)
            import io
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_content = file_stream.getvalue()
        
        escritorio = await crud.get_escritorio(db, current_user.escritorio_id)
        storage = get_storage_provider(escritorio)
        
        safe_title = slugify(request.titulo_documento)
        unique_filename = f"{safe_title}_{uuid.uuid4().hex[:6]}.docx"
        
        # Organização: cliente_{id}/documentos ou escritorio/documentos
        if request.cliente_id:
            relative_dir = f"cliente_{request.cliente_id}/documentos"
        else:
            relative_dir = "escritorio/documentos"
            
        db_path = await storage.save_file(file_content, relative_dir, unique_filename)
        
    except Exception as e:
        import traceback
        print(f"ERRO CRÍTICO NO REDATOR: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro ao processar template .docx: {str(e)}")

    doc_create = schemas.DocumentoClienteCreate(nome=request.titulo_documento, cliente_id=request.cliente_id)
    db_doc = await crud.create_documento_cliente(
        db, 
        doc_create, 
        db_path, 
        current_user.escritorio_id
    )

    return db_doc
