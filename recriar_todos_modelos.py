
import os
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

PASTA_MODELOS = "d:/Projetos/advtools/modelos_padrao"

def limpar_pasta():
    if os.path.exists(PASTA_MODELOS):
        shutil.rmtree(PASTA_MODELOS)
    os.makedirs(PASTA_MODELOS)
    print("Pasta de modelos limpa.")

def criar_modelo_base(nome_arquivo, titulo, tipo='geral'):
    doc = Document()
    
    # Cabeçalho Padrão
    heading = doc.add_heading(titulo.upper(), 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # Qualificação (Cliente)
    p = doc.add_paragraph()
    p.add_run('EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA ____ VARA ____________________ DA COMARCA DE ____________________').bold = True
    doc.add_paragraph('')
    
    p = doc.add_paragraph()
    run_nome = p.add_run('{{ cliente_nome }}')
    run_nome.bold = True
    p.add_run(', inscrito(a) no CPF/CNPJ sob nº ')
    p.add_run('{{ cliente_doc }}').bold = True
    p.add_run(', residente e domiciliado(a) em ')
    p.add_run('{{ cliente_endereco }}')
    p.add_run(', por seu advogado infra-assinado, vem respeitosamente à presença de Vossa Excelência propor:')
    
    doc.add_paragraph('')
    p_titulo = doc.add_paragraph(titulo.upper())
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.style = 'Title'
    doc.add_paragraph('')

    # Corpo
    doc.add_heading('DOS FATOS', level=1)
    if tipo == 'peticao':
        doc.add_paragraph('{{ fatos }}') # Tag usada na rota novo_documento
    elif tipo == 'contrato':
        doc.add_paragraph('{{ conteudo_ia }}')
    else:
        doc.add_paragraph('{{ conteudo_ia }}') # Fallback
        
    doc.add_paragraph('')
    
    doc.add_heading('DO DIREITO', level=1)
    doc.add_paragraph('(Fundamentação jurídica gerada pela IA ou inserida manualmente)')

    doc.add_heading('DOS PEDIDOS', level=1)
    doc.add_paragraph('Diante do exposto, requer:')
    doc.add_paragraph('a) A procedência total da ação;')
    doc.add_paragraph('b) A citação do requerido;')
    doc.add_paragraph('c) A condenação em custas e honorários.')
    doc.add_paragraph('')
    
    doc.add_paragraph('Dá-se à causa o valor de R$ ________________.')
    
    doc.add_paragraph('')
    doc.add_paragraph('Nestes termos,')
    doc.add_paragraph('Pede deferimento.')
    
    doc.add_paragraph('')
    p_data = doc.add_paragraph('Local, {{ data_hoje }}')
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph('')
    doc.add_paragraph('____________________________________')
    doc.add_paragraph('ADVOGADO OAB/UF')

    doc.save(os.path.join(PASTA_MODELOS, nome_arquivo))
    print(f"Criado: {nome_arquivo}")

def criar_contrato_honorarios():
    doc = Document()
    heading = doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS JURÍDICOS', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.add_run('CONTRATANTE: ').bold = True
    p.add_run('{{ cliente_nome }}').bold = True
    p.add_run(', CPF: {{ cliente_doc }}, Endereço: {{ cliente_endereco }}.')
    doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.add_run('CONTRATADO: ').bold = True
    p.add_run('SEU ESCRITÓRIO DE ADVOCACIA...')
    doc.add_paragraph('')

    doc.add_heading('CLÁUSULA 1 - DO OBJETO', level=1)
    doc.add_paragraph('O presente contrato tem por objeto: {{ servico_tipo }}')
    doc.add_paragraph('{{ conteudo_ia }}') # Detalhes da IA
    
    doc.add_heading('CLÁUSULA 2 - DO VALOR', level=1)
    doc.add_paragraph('Valor Total: R$ {{ valor_total }}')
    doc.add_paragraph('Forma de Pagamento: {{ forma_pagamento }}')
    doc.add_paragraph('Detalhes: {{ detalhes_pagamento }}')
    
    doc.add_paragraph('')
    doc.add_paragraph('Local, {{ data_hoje }}')
    
    doc.add_paragraph('')
    doc.add_paragraph('__________________________')
    doc.add_paragraph('CONTRATANTE')
    
    doc.add_paragraph('')
    doc.add_paragraph('__________________________')
    doc.add_paragraph('CONTRATADO')

    doc.save(os.path.join(PASTA_MODELOS, 'Modelo_Contrato_Honorarios.docx'))
    print("Criado: Modelo_Contrato_Honorarios.docx")

def criar_procuracao():
    doc = Document()
    heading = doc.add_heading('PROCURAÇÃO AD JUDICIA', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.add_run('OUTORGANTE: ').bold = True
    p.add_run('{{ cliente_nome }}').bold = True
    p.add_run(', nacionalidade..., estado civil..., profissão..., inscrito no CPF sob nº {{ cliente_doc }}, residente e domiciliado em {{ cliente_endereco }}.')
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('OUTORGADO: ').bold = True
    p.add_run('ADVOGADO..., inscrito na OAB/UF sob nº...')
    
    doc.add_paragraph('')
    doc.add_paragraph('PODERES: Pelo presente instrumento, o outorgante confere ao outorgado amplos poderes para o foro em geral, com cláusula "ad judicia et extra", para propor ou defender seus interesses em qualquer juízo, instância ou tribunal.')
    
    doc.add_paragraph('')
    doc.add_paragraph('Local, {{ data_hoje }}')
    
    doc.add_paragraph('')
    doc.add_paragraph('__________________________')
    doc.add_paragraph('{{ cliente_nome }}')
    
    doc.save(os.path.join(PASTA_MODELOS, 'Modelo_Procuracao.docx'))
    print("Criado: Modelo_Procuracao.docx")

if __name__ == "__main__":
    limpar_pasta()
    
    # Criar Modelos Cíveis
    criar_modelo_base('Peticao_Inicial_Civel.docx', 'Petição Inicial', 'peticao')
    criar_modelo_base('Contestacao_Civel.docx', 'Contestação', 'peticao')
    criar_modelo_base('Agravo_Instrumento.docx', 'Agravo de Instrumento', 'peticao')
    criar_modelo_base('Apelacao.docx', 'Apelação', 'peticao')
    
    # Criar Modelos Trabalhistas
    criar_modelo_base('Reclamacao_Trabalhista.docx', 'Reclamação Trabalhista', 'peticao')
    criar_modelo_base('Contestacao_Trabalhista.docx', 'Contestação Trabalhista', 'peticao')
    criar_modelo_base('Recurso_Ordinario.docx', 'Recurso Ordinário', 'peticao')
    
    # Criar Modelos Criminais
    criar_modelo_base('Habeas_Corpus.docx', 'Habeas Corpus', 'peticao')
    criar_modelo_base('Resposta_Acusacao.docx', 'Resposta à Acusação', 'peticao')
    criar_modelo_base('Liberdade_Provisoria.docx', 'Pedido de Liberdade Provisória', 'peticao')
    
    # Outros
    criar_modelo_base('Mandado_Seguranca.docx', 'Mandado de Segurança', 'peticao')
    criar_contrato_honorarios()
    criar_procuracao()
    
    print("Todos os modelos foram recriados e padronizados com sucesso!")
