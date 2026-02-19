from docx import Document

import os

def criar_modelo_ia():
    document = Document()
    document.add_heading('PROCURAÇÃO AD JUDICIA', 0)

    # Texto Padrão
    p = document.add_paragraph('OUTORGANTE: ')
    p.add_run('{{ nome_cliente }}').bold = True
    p.add_run(', {{ nacionalidade }}, {{ estado_civil }}, portador do RG nº {{ rg }} e CPF nº {{ cpf }}, residente em {{ endereco }}.')

    document.add_paragraph('\nOUTORGADO: FERNANDO COZAC, advogado inscrito na OAB/GO...')
    document.add_paragraph('\nPODERES GERAIS: Os da cláusula ad judicia et extra para o foro em geral.')

    # --- O CAMPO MÁGICO DA IA ---
    document.add_heading('PODERES ESPECÍFICOS (IA):', level=2)
    p_ia = document.add_paragraph('{{ poderes_ia }}') 
    p_ia.italic = True # Destaca o texto gerado

    document.add_paragraph('\nCidade: {{ cidade }}, Data: {{ data_hoje }}')
    document.add_paragraph('\n____________________________________\nAssinatura do Outorgante')

    nome_final = os.path.join("modelos", "modelo_procuracao.docx")
    document.save(nome_final)
    print(f"✅ Modelo Word atualizado com campo '{{ poderes_ia }}'!")

if __name__ == "__main__":
    criar_modelo_ia()