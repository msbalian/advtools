
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_contract_template():
    doc = Document()
    
    # Título
    heading = doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS JURÍDICOS', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('') # Espaço

    # Qualificação
    p_qualif = doc.add_paragraph()
    p_qualif.add_run('Pelo presente instrumento particular, de um lado ').bold = True
    p_qualif.add_run('DOUTOR FULANO DE TAL, advogado inscrito na OAB/UF sob nº 00.000, com escritório em..., doravante denominado ')
    p_qualif.add_run('CONTRATADO(A)').bold = True
    p_qualif.add_run(', e de outro lado ')
    
    # TAGS DO CLIENTE
    run_cliente = p_qualif.add_run('{{ cliente_nome }}')
    run_cliente.bold = True
    run_cliente.underline = True
    
    p_qualif.add_run(', inscrito(a) no CPF/CNPJ sob nº ')
    p_qualif.add_run('{{ cliente_doc }}').bold = True
    p_qualif.add_run(', residente e domiciliado(a) em ')
    p_qualif.add_run('{{ cliente_endereco }}')
    p_qualif.add_run(', doravante denominado(a) ')
    p_qualif.add_run('CONTRATANTE').bold = True
    p_qualif.add_run(', celebram o presente contrato mediante as cláusulas abaixo:')

    doc.add_paragraph('')

    # Cláusula 1 - Objeto (Aqui entra a IA)
    doc.add_heading('CLÁUSULA PRIMEIRA - DO OBJETO E NATUREZA DOS SERVIÇOS', level=1)
    
    p_obj = doc.add_paragraph()
    p_obj.add_run('O presente contrato tem por objeto a prestação de serviços advocatícios para: ')
    p_obj.add_run('{{ servico_tipo }}').bold = True
    p_obj.add_run('. ')
    
    # TAG DA IA (Vai injetar texto longo)
    doc.add_paragraph('{{ conteudo_ia }}')
    
    doc.add_paragraph('')

    # Cláusula 2 - Honorários
    doc.add_heading('CLÁUSULA SEGUNDA - DOS HONORÁRIOS', level=1)
    
    p_hon = doc.add_paragraph('Pelos serviços prestados, o(a) CONTRATANTE pagará ao(à) CONTRATADO(A) o valor total de: ')
    p_hon.add_run('R$ {{ valor_total }}').bold = True
    p_hon.add_run('.')
    
    doc.add_paragraph('Forma de Pagamento: {{ forma_pagamento }}')
    doc.add_paragraph('Detalhes Adicionais: {{ detalhes_pagamento }}')

    doc.add_paragraph('')
    
    # Cláusula 3 - Foro
    doc.add_heading('CLÁUSULA TERCEIRA - DO FORO', level=1)
    doc.add_paragraph('Fica eleito o foro da Comarca de Local/UF para dirimir quaisquer dúvidas oriundas deste contrato.')
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Assinaturas
    p_ass = doc.add_paragraph('Local/UF, {{ data_hoje }}')
    p_ass.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    p_adv = doc.add_paragraph('__________________________________________')
    p_adv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_adv.add_run('\nADVOGADO(A) CONTRATADO(A)')

    doc.add_paragraph('')
    doc.add_paragraph('')
    
    p_cli = doc.add_paragraph('__________________________________________')
    p_cli.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cli.add_run('\n{{ cliente_nome }}\nCONTRATANTE')

    # Salva
    doc.save('d:/Projetos/advtools/modelos_padrao/Modelo_Contrato_Correto.docx')
    print("Modelo criado com sucesso: Modelo_Contrato_Correto.docx")

if __name__ == "__main__":
    try:
        create_contract_template()
    except Exception as e:
        print(f"Erro: {e}")
